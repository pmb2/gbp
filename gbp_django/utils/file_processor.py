import os
import magic
import docx
import PyPDF2
import io
import json
import uuid
import time
from typing import Dict, Any, List, Optional
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from ..models import KnowledgeFile, Business, KnowledgeChunk
from .embeddings import generate_embedding

def get_file_mime_type(file_content: bytes) -> str:
    """Determine file type using python-magic"""
    mime = magic.Magic(mime=True)
    return mime.from_buffer(file_content)

def process_text_file(content: bytes) -> str:
    """Process plain text files"""
    return content.decode('utf-8', errors='ignore')

def process_docx(content: bytes) -> str:
    """Process .docx files with enhanced error handling"""
    import io
    try:
        doc = docx.Document(io.BytesIO(content))
        
        # Process paragraphs with better formatting
        text_parts = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:  # Only add non-empty paragraphs
                text_parts.append(text)
                
        # Process tables if present
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
                    
        # Join with double newlines for better separation
        processed_text = '\n\n'.join(text_parts)
        
        if not processed_text.strip():
            raise ValueError("No readable text content found in DOCX file")
            
        print(f"Successfully processed DOCX content:")
        print(f"Found {len(text_parts)} content blocks")
        print(f"Total text length: {len(processed_text)} characters")
        
        return processed_text
        
    except Exception as e:
        error_msg = f"DOCX processing failed: {str(e)}\n"
        error_msg += "This might be due to corrupt file format or unsupported DOCX features"
        raise ValueError(error_msg)

def process_pdf(content: bytes) -> str:
    """Process PDF files"""
    pdf_file = io.BytesIO(content)
    reader = PyPDF2.PdfReader(pdf_file)
    text = ""
    for page in reader.pages:
        text += page.extract_text() + "\n"
    return text

def process_markdown(content: bytes) -> str:
    """Process markdown files"""
    return content.decode('utf-8', errors='ignore')

import logging

logger = logging.getLogger(__name__)

def store_file_content(business_id: str, file_obj: Any, filename: str) -> Dict[str, Any]:
    """Store file content and generate embeddings with extensive error handling and chunking"""
    print(f"[INFO] Starting file processing for {filename} (Business ID: {business_id})")
    file_id = str(uuid.uuid4())  # Generate file ID once per file
    try:
        # Validate file size (10MB limit)
        logger.debug(f"Original filename: {filename}")
        MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
        file_obj.seek(0, 2)  # Seek to end
        file_size = file_obj.tell()
        file_obj.seek(0)  # Reset to beginning
        
        if file_size > MAX_FILE_SIZE:
            raise ValueError(f"File size exceeds limit of {MAX_FILE_SIZE/1024/1024}MB")

        # Read file content with timeout
        try:
            content = file_obj.read()
        except Exception as e:
            raise IOError(f"Failed to read file content: {str(e)}")

        # Validate file type
        try:
            mime_type = get_file_mime_type(content)
        except Exception as e:
            raise ValueError(f"Failed to determine file type: {str(e)}")

        # Truncate the MIME type if it exceeds the max_length
        max_mime_length = KnowledgeFile._meta.get_field('file_type').max_length
        if len(mime_type) > max_mime_length:
            mime_type = mime_type[:max_mime_length]
        try:
            print(f"\nProcessing file content:")
            print(f"[INFO] MIME type: {mime_type}")
            print(f"[INFO] File size: {file_size/1024:.1f}KB")
            
            if mime_type == 'text/plain':
                text_content = process_text_file(content)
            elif mime_type == 'application/pdf':
                text_content = process_pdf(content)
            elif mime_type in [
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'application/msword'
            ]:
                text_content = process_docx(content)
            elif mime_type == 'text/markdown':
                text_content = process_markdown(content)
            else:
                raise ValueError(f"Unsupported file type: {mime_type}")
                
            print(f"Successfully extracted text content")
            print(f"Extracted text length: {len(text_content)} characters")
            print(f"Text preview: {text_content[:200]}...")
            
        except Exception as e:
            error_msg = f"Failed to process file content: {str(e)}\n"
            error_msg += f"File: {filename}\n"
            error_msg += f"MIME type: {mime_type}\n"
            error_msg += f"Size: {file_size/1024:.1f}KB"
            print(f"ERROR: {error_msg}")
            raise ValueError(error_msg)

        # Validate text content
        if not text_content or len(text_content.strip()) == 0:
            raise ValueError("Extracted text content is empty")

        # Generate embedding with validation
        try:
            print("\nStarting content preprocessing...")
            
            # Clean and normalize text content
            if isinstance(text_content, str):
                text_content = text_content.replace('\r', '\n')
                text_content = text_content.replace('\t', ' ')
                text_content = '\n'.join(line.strip() for line in text_content.split('\n'))
                text_content = '\n'.join(filter(None, text_content.split('\n')))  # Remove empty lines
            else:
                raise ValueError("Processed content is not a string")
            
            if not text_content.strip():
                raise ValueError("Text content is empty after cleaning")
                    
            print(f"Cleaned text content length: {len(text_content)} characters")
                
            # Enhanced content chunking with optimized sizes for RAG
            MAX_CHUNK_SIZE = 2000  # Optimized for LLM context window
            MIN_CHUNK_SIZE = 200   # Ensure meaningful semantic chunks
            OVERLAP_SIZE = 100     # Add overlap between chunks for context continuity
            chunks = []
            
            # Enhanced text preprocessing
            text_content = text_content.replace('\r\n', '\n').replace('\r', '\n')
            
            # Split on multiple newlines while preserving structure
            paragraphs = []
            current_para = []
            
            for line in text_content.split('\n'):
                line = line.strip()
                if line:
                    current_para.append(line)
                elif current_para:
                    paragraphs.append(' '.join(current_para))
                    current_para = []
                    
            if current_para:
                paragraphs.append(' '.join(current_para))
                
            # Filter out any remaining empty paragraphs
            paragraphs = [p for p in paragraphs if p.strip()]
            
            print(f"\nStarting content chunking:")
            print(f"Total content length: {len(text_content)} characters")
            print(f"Found {len(paragraphs)} paragraphs")
            
            # Improved chunking with overlap and semantic boundaries
            current_chunk = []
            current_length = 0
            last_overlap = ""
            
            for para in paragraphs:
                para_length = len(para)
                
                # Check if adding this paragraph would exceed max size
                if current_length + para_length > MAX_CHUNK_SIZE:
                    if current_chunk and sum(len(p) for p in current_chunk) >= MIN_CHUNK_SIZE:
                        # Add overlap from previous chunk if available
                        if last_overlap:
                            current_chunk.insert(0, last_overlap)
                        
                        chunk_text = ' '.join(current_chunk)
                        chunks.append(chunk_text)
                        print(f"Created chunk of {len(chunk_text)} characters")
                        
                        # Store last part of current chunk for overlap
                        overlap_text = ' '.join(current_chunk[-2:]) if len(current_chunk) > 1 else current_chunk[-1]
                        last_overlap = overlap_text[-OVERLAP_SIZE:] if len(overlap_text) > OVERLAP_SIZE else overlap_text
                        
                    # Start new chunk with overlap
                    current_chunk = [last_overlap, para] if last_overlap else [para]
                    current_length = len(last_overlap) + para_length if last_overlap else para_length
                else:
                    current_chunk.append(para)
                    current_length += para_length
            
            # Add final chunk if it meets minimum size
            if current_chunk:
                final_chunk = ' '.join(current_chunk)
                if len(final_chunk) >= MIN_CHUNK_SIZE:
                    chunks.append(final_chunk)
                    print(f"Created final chunk of {len(final_chunk)} characters")
            
            print(f"Created {len(chunks)} chunks for processing")
            
            # Generate embeddings with improved chunking and retries
            embeddings = []
            max_retries = 3
            embedding_vector = None  # Initialize embedding_vector

            for idx, chunk in enumerate(chunks):
                retry_count = 0
                while retry_count < max_retries:
                    try:
                        print(f"\nProcessing chunk {idx + 1}/{len(chunks)}")
                        print(f"Chunk length: {len(chunk)} characters")
                        print(f"Chunk preview: {chunk[:100]}...")
                        
                        embedding = generate_embedding(chunk)
                        
                        if embedding is None:
                            print(f"Warning: Embedding generation returned None for chunk {idx + 1}")
                            print(f"Chunk content: {chunk[:500]}...")
                            print("Retrying with cleaned content...")
                            # Try cleaning the content and retry
                            cleaned_chunk = ' '.join(chunk.split())  # Remove extra whitespace
                            embedding = generate_embedding(cleaned_chunk)
                            if embedding is None:
                                print("Retry failed - skipping chunk")
                                continue
                        
                        
                        embeddings.append({
                            'text': chunk,
                            'embedding': embedding
                        })
                        embedding_vector = embedding  # Assign the last successful embedding
                        print(f"Successfully generated embedding for chunk {idx + 1}")
                        break  # Exit retry loop on success
                        
                    except Exception as e:
                        print(f"Error processing chunk {idx + 1}: {str(e)}")
                        retry_count += 1
                        if retry_count >= max_retries:
                            print(f"Max retries reached for chunk {idx + 1}")
                            break
                        continue
            
            if not embeddings:
                error_msg = "Failed to generate embeddings.\n"
                error_msg += f"Processed {len(chunks)} chunks but none were successful.\n"
                error_msg += f"Original text length: {len(text_content)} chars\n"
                error_msg += f"Number of paragraphs: {len(paragraphs)}\n"
                error_msg += f"First 500 chars of content: {text_content[:500]}\n"
                error_msg += f"Chunk sizes: {[len(c) for c in chunks]}\n"
                error_msg += f"MIME type: {mime_type}\n"
                error_msg += f"File size: {file_size/1024:.1f}KB\n"
                error_msg += "Check the logs for specific chunk processing errors."
                print(error_msg)
                
                # Try one final time with the entire content as a single chunk
                if len(text_content) < 4000:  # Only try if content is reasonably sized
                    print("Attempting to process entire content as single chunk...")
                    embedding = generate_embedding(text_content)
                    if embedding:
                        embeddings.append({
                            'text': text_content,
                            'embedding': embedding
                        })
                        print("Successfully generated embedding for entire content")
                    else:
                        raise ValueError("Document processing failed - unable to generate embeddings")
                else:
                    raise ValueError("Document processing failed - content too large for single chunk processing")
        except Exception as e:
            raise ValueError(f"Embedding generation failed: {str(e)}")

        # Verify business exists
        try:
            business = Business.objects.get(business_id=business_id)
        except Business.DoesNotExist:
            raise ValueError(f"Business with ID {business_id} not found")

        # Store file safely
        try:
            file_path = f'knowledge_base/{business_id}/{filename}'
            logger.debug(f"Storing file at path: {file_path}")
            saved_path = default_storage.save(file_path, ContentFile(content))
            logger.info(f"File stored successfully at: {saved_path}")
        except Exception as e:
            raise IOError(f"Failed to store file: {str(e)}")

        # Create KnowledgeFile instance (without the embedding field)
        knowledge_file = KnowledgeFile.objects.create(
            business=business,
            file_name=filename,
            file_path=saved_path,
            file_type=mime_type,
            file_size=file_size,
            content=text_content  # Consider removing this field if content is too large
        )

        # Save each chunk and its embedding as a KnowledgeChunk
        for idx, chunk_data in enumerate(embeddings):
            embedding = chunk_data['embedding']
            if embedding is None or len(embedding) != 1536:
                print(f"[ERROR] Invalid embedding for chunk {idx} in file {filename}")
                continue  # Skip this chunk

            print(f"[DEBUG] Valid embedding generated for chunk {idx} in file {filename}")

            KnowledgeChunk.objects.create(
                knowledge_file=knowledge_file,
                business=business,  # Set the business
                content=chunk_data['text'],
                embedding=embedding,
                position=idx
            )

        # Return file info
        print(f"[DEBUG] KnowledgeFile ID: {knowledge_file.id}")
        return {
            'id': knowledge_file.id,
            'name': filename,
            'size': file_size,
            'type': mime_type,
            'path': saved_path,
        }

    except (ValueError, IOError) as e:
        # Log specific error types
        print(f"Error processing file {filename}: {str(e)}")
        raise
    except Exception as e:
        # Log unexpected errors
        print(f"Unexpected error processing file {filename}: {str(e)}")
        raise ValueError(f"Unexpected error: {str(e)}")

def process_folder(business_id: str, folder_path: str) -> List[Dict[str, Any]]:
    """Process all files in a folder"""
    results = []
    allowed_extensions = {'.txt', '.pdf', '.docx', '.md'}

    for root, _, files in os.walk(folder_path):
        for filename in files:
            if os.path.splitext(filename)[1].lower() in allowed_extensions:
                file_path = os.path.join(root, filename)
                with open(file_path, 'rb') as f:
                    try:
                        result = store_file_content(business_id, f, filename)
                        results.append(result)
                    except Exception as e:
                        print(f"Error processing {filename}: {str(e)}")
                        continue

    return results
